from docx import Document


def format_author(author_str):
    author_formatted_str = ""
    authors_arr = author_str.split(" and ")
    count = 0
    for author in authors_arr:
        author_lastname = author.split(", ")[0]
        try:
            author_forname = author.split(", ")[1]
            if len(author_lastname) == 1:
                helper = author_lastname
                author_lastname = author_forname
                author_forname = helper
            author_forname_initial = author_forname[0]
        except:
            author_forname_initial = ""

        if count == 0:
            author_formatted_str = author_lastname + " " + author_forname_initial
            count = count + 1
        else:
            author_formatted_str = (
                author_formatted_str
                + ", "
                + author_lastname
                + " "
                + author_forname_initial
            )

    return author_formatted_str


def parseBibTexToString(file, from_year):
    import bibtexparser

    # library = bibtexparser.parse_file(file)
    library = bibtexparser.parse_string(file)

    print(
        f"Parsed {len(library.blocks)} blocks, including:"
        f"\n\t{len(library.entries)} entries"
        f"\n\t{len(library.comments)} comments"
        f"\n\t{len(library.strings)} strings and"
        f"\n\t{len(library.preambles)} preambles"
    )

    results_arr = []
    for el in library.entries:
        entry = el.fields_dict
        year = entry["year"].value
        author = format_author(entry["author"].value)
        title = entry["title"].value
        title = title.replace("{", "")
        title = title.replace("}", "")
        journal = ""
        pages = ""
        volume = ""

        try:
            pages = " " + entry["pages"].value + "."
        except:
            print("Pages not available")
        try:
            if not pages:
                volume = " " + entry["volume"].value + "."
            else:
                volume = " " + entry["volume"].value + ": "
        except:
            print("Volume not available")

        try:
            journal = " " + entry["journal"].value
            journal = journal.replace("{", "")
            journal = journal.replace("}", "")
        except:
            print("Journal not available")
            volume = ""

        entryString = f"""{author} ({year}) {title}. {journal}{volume}{pages}"""

        if int(year) > from_year - 1:
            results_arr.append(entryString)

    if len(results_arr) == 0:
        results_arr.append(f"No citations for year {from_year}")
    return results_arr


def parseBibTexToDocx(file):
    import bibtexparser

    library = bibtexparser.parse_file(file)

    document = Document()

    document.add_heading("Citations", 0)

    for el in library.entries:
        entry = el.fields_dict
        author = format_author(entry["author"].value)
        year = entry["year"].value
        if int(year) < 2022:
            continue

        title = entry["title"].value
        title = title.replace("{", "")
        title = title.replace("}", "")
        journal = entry["journal"].value
        try:
            volume = entry["volume"].value
            pages = entry["pages"].value
        except:
            print("Volume or pages not available")

        p = document.add_paragraph("")
        p.add_run(author)
        p.add_run(" (" + year + ") " + title + ". ")
        p.add_run(journal).italic = True
        p.add_run(": " + pages + ".\n")
        c = document.add_paragraph("")

    document.save("citations.docx")
    return "document saved"
